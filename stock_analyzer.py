import streamlit as st
import yfinance as yf
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from datetime import datetime, timedelta
import io
import re

try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

st.set_page_config(
    page_title="Stock Analyzer Pro",
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Clean, Modern Light Theme CSS
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@300;400;500;600;700;800&display=swap');

    * {
        font-family: 'Plus Jakarta Sans', sans-serif;
    }

    .stApp {
        background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
    }

    #MainMenu, footer, header {visibility: hidden;}

    /* Main Title - Extra Large */
    .main-title {
        font-size: 4.5rem;
        font-weight: 800;
        text-align: center;
        color: #0f172a;
        margin-bottom: 0.3rem;
        letter-spacing: -2px;
        line-height: 1.1;
    }

    .main-title span {
        background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }

    .sub-title {
        text-align: center;
        color: #64748b;
        font-size: 1.4rem;
        font-weight: 500;
        margin-bottom: 3rem;
    }

    /* Metric Cards - Clean & Bold */
    .metric-card {
        background: #ffffff;
        padding: 35px;
        border-radius: 24px;
        text-align: center;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.04);
        border: 1px solid #e2e8f0;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        margin-bottom: 1.5rem;
    }

    .metric-card:hover {
        transform: translateY(-8px);
        box-shadow: 0 20px 40px rgba(0, 0, 0, 0.08);
    }

    .metric-card-price { border-top: 5px solid #3b82f6; }
    .metric-card-ma { border-top: 5px solid #8b5cf6; }
    .metric-card-atr { border-top: 5px solid #f59e0b; }

    .metric-label {
        font-size: 1rem;
        font-weight: 600;
        color: #94a3b8;
        text-transform: uppercase;
        letter-spacing: 2px;
        margin-bottom: 15px;
    }

    .metric-value {
        font-size: 3.8rem;
        font-weight: 800;
        margin: 10px 0;
        letter-spacing: -2px;
    }

    .metric-value-price { color: #3b82f6; }
    .metric-value-ma { color: #8b5cf6; }
    .metric-value-atr { color: #f59e0b; }

    .metric-change {
        font-size: 1.2rem;
        font-weight: 700;
        padding: 10px 20px;
        border-radius: 50px;
        display: inline-block;
        margin-top: 12px;
    }

    .metric-change-up {
        background: #dcfce7;
        color: #16a34a;
    }

    .metric-change-down {
        background: #fee2e2;
        color: #dc2626;
    }

    /* Company Name */
    .company-name {
        font-size: 2.5rem;
        font-weight: 800;
        color: #0f172a;
        text-align: center;
        margin: 2rem 0;
        padding: 1.5rem;
        background: #ffffff;
        border-radius: 20px;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.04);
    }

    /* Text Input - Large & Clean */
    .stTextInput > div > div > input {
        background: #ffffff !important;
        border: 3px solid #e2e8f0 !important;
        border-radius: 20px !important;
        color: #0f172a !important;
        font-size: 1.5rem !important;
        padding: 25px 30px !important;
        text-align: center !important;
        font-weight: 700 !important;
        letter-spacing: 3px !important;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.04) !important;
    }

    .stTextInput > div > div > input:focus {
        border-color: #3b82f6 !important;
        box-shadow: 0 0 0 4px rgba(59, 130, 246, 0.15) !important;
    }

    .stTextInput > div > div > input::placeholder {
        color: #94a3b8 !important;
        font-weight: 500 !important;
        letter-spacing: 1px !important;
    }

    /* Tabs - Modern Pill Style */
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px;
        background: #ffffff;
        padding: 10px;
        border-radius: 20px;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.04);
        border: 1px solid #e2e8f0;
    }

    .stTabs [data-baseweb="tab"] {
        background: transparent;
        border-radius: 14px;
        color: #64748b;
        font-weight: 700;
        font-size: 1.2rem;
        padding: 16px 32px;
    }

    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%) !important;
        color: white !important;
    }

    /* Buttons - Bold & Vibrant */
    .stButton > button {
        background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 16px !important;
        padding: 20px 40px !important;
        font-weight: 800 !important;
        font-size: 1.3rem !important;
        letter-spacing: 0.5px !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 15px rgba(59, 130, 246, 0.3) !important;
    }

    .stButton > button:hover {
        transform: translateY(-3px) !important;
        box-shadow: 0 8px 25px rgba(59, 130, 246, 0.4) !important;
    }

    /* File Uploader */
    .stFileUploader > div {
        background: #ffffff !important;
        border: 3px dashed #cbd5e1 !important;
        border-radius: 24px !important;
        padding: 50px !important;
    }

    /* DataFrames */
    .stDataFrame {
        background: #ffffff !important;
        border-radius: 20px !important;
        border: 1px solid #e2e8f0 !important;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.04) !important;
    }

    /* Alerts */
    .stAlert {
        background: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
        border-radius: 16px !important;
        color: #475569 !important;
        font-size: 1.1rem !important;
    }

    /* Expander */
    .streamlit-expanderHeader {
        background: #ffffff !important;
        border-radius: 16px !important;
        font-weight: 700 !important;
        font-size: 1.2rem !important;
        color: #0f172a !important;
    }

    /* Popular Tickers - Chips */
    .ticker-chip {
        background: linear-gradient(135deg, rgba(59, 130, 246, 0.1) 0%, rgba(139, 92, 246, 0.1) 100%);
        border: 2px solid rgba(59, 130, 246, 0.2);
        border-radius: 14px;
        padding: 18px 30px;
        text-align: center;
        font-weight: 800;
        font-size: 1.4rem;
        color: #3b82f6;
        letter-spacing: 2px;
    }

    /* Footer */
    .footer {
        text-align: center;
        color: #94a3b8;
        font-size: 1rem;
        padding: 3rem 0;
        margin-top: 4rem;
        border-top: 1px solid #e2e8f0;
        font-weight: 500;
    }

    /* Section Headers */
    .section-header {
        font-size: 2rem;
        font-weight: 800;
        color: #0f172a;
        margin: 3rem 0 1.5rem 0;
        padding-bottom: 0.8rem;
        border-bottom: 3px solid #e2e8f0;
    }

    /* Download Button */
    .stDownloadButton > button {
        background: #ffffff !important;
        border: 3px solid #3b82f6 !important;
        color: #3b82f6 !important;
        font-weight: 700 !important;
    }

    .stDownloadButton > button:hover {
        background: rgba(59, 130, 246, 0.1) !important;
    }

    /* Progress Bar */
    .stProgress > div > div {
        background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%) !important;
    }

    /* Text Area */
    .stTextArea > div > div > textarea {
        background: #ffffff !important;
        border: 2px solid #e2e8f0 !important;
        border-radius: 16px !important;
        color: #0f172a !important;
        font-size: 1.1rem !important;
        font-weight: 600 !important;
    }

    /* Metrics */
    [data-testid="stMetricValue"] {
        font-size: 1.8rem !important;
        font-weight: 800 !important;
        color: #0f172a !important;
    }

    [data-testid="stMetricLabel"] {
        font-size: 1rem !important;
        font-weight: 600 !important;
        color: #64748b !important;
    }

    /* Summary Cards */
    .summary-card {
        background: #ffffff;
        border: 1px solid #e2e8f0;
        border-radius: 20px;
        padding: 30px;
        text-align: center;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.04);
    }

    .summary-number {
        font-size: 3.5rem;
        font-weight: 800;
        background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }

    .summary-label {
        color: #64748b;
        font-size: 1.2rem;
        font-weight: 600;
        margin-top: 10px;
    }

    /* Upload Area */
    .upload-area {
        text-align: center;
        padding: 60px;
        color: #64748b;
        background: #ffffff;
        border-radius: 24px;
        border: 3px dashed #cbd5e1;
    }

    .upload-icon {
        font-size: 4rem;
        margin-bottom: 1rem;
    }

    .upload-text {
        font-size: 1.4rem;
        font-weight: 600;
        color: #475569;
    }

    .upload-subtext {
        font-size: 1.1rem;
        margin-top: 0.5rem;
        color: #94a3b8;
    }
</style>
""", unsafe_allow_html=True)

def calculate_atr(df, period=14):
    high = df['High']
    low = df['Low']
    close = df['Close']
    tr1 = high - low
    tr2 = abs(high - close.shift(1))
    tr3 = abs(low - close.shift(1))
    tr = pd.concat([tr1, tr2, tr3], axis=1).max(axis=1)
    atr = tr.rolling(window=period).mean()
    return atr

def get_stock_data(ticker, days=400):
    end_date = datetime.now()
    start_date = end_date - timedelta(days=days)
    stock = yf.Ticker(ticker)
    df = stock.history(start=start_date, end=end_date)
    if df.empty:
        return None, None
    df['MA150'] = df['Close'].rolling(window=150).mean()
    df['ATR'] = calculate_atr(df, period=14)
    return df, stock.info

def get_stock_summary(ticker):
    try:
        df, info = get_stock_data(ticker)
        if df is None or df.empty:
            return None
        current_price = df['Close'].iloc[-1]
        ma150 = df['MA150'].iloc[-1]
        atr = df['ATR'].iloc[-1]
        if pd.notna(ma150) and ma150 != 0:
            gap_pct = ((current_price - ma150) / ma150) * 100
        else:
            gap_pct = None
        return {
            'Ticker': ticker,
            'Current Price': current_price,
            '150-Day MA': ma150,
            'Gap %': gap_pct,
            'ATR (14)': atr
        }
    except Exception as e:
        return None

def extract_tickers_from_text(text):
    common_words = {'A', 'I', 'AM', 'AN', 'AS', 'AT', 'BE', 'BY', 'DO', 'GO', 'HE', 'IF',
                   'IN', 'IS', 'IT', 'ME', 'MY', 'NO', 'OF', 'OK', 'ON', 'OR', 'SO', 'TO',
                   'UP', 'US', 'WE', 'THE', 'AND', 'FOR', 'ARE', 'BUT', 'NOT', 'YOU', 'ALL',
                   'CAN', 'HAD', 'HER', 'WAS', 'ONE', 'OUR', 'OUT', 'PDF', 'USD', 'EUR'}
    potential_tickers = re.findall(r'\b[A-Z]{1,5}\b', text.upper())
    tickers = [t for t in potential_tickers if t not in common_words]
    seen = set()
    unique_tickers = []
    for t in tickers:
        if t not in seen:
            seen.add(t)
            unique_tickers.append(t)
    return unique_tickers

def read_excel_tickers(file):
    try:
        df = pd.read_excel(file)
        tickers = []
        for col in df.columns:
            col_values = df[col].dropna().astype(str).tolist()
            for val in col_values:
                extracted = extract_tickers_from_text(val)
                tickers.extend(extracted)
        return list(dict.fromkeys(tickers))
    except Exception as e:
        st.error(f"Error reading Excel file: {e}")
        return []

def read_pdf_tickers(file):
    if not PDF_SUPPORT:
        st.error("PDF support not available.")
        return []
    try:
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return extract_tickers_from_text(text)
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return []

def read_word_tickers(file):
    if not DOCX_SUPPORT:
        st.error("Word support not available.")
        return []
    try:
        doc = Document(file)
        text = " ".join([para.text for para in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += " " + cell.text
        return extract_tickers_from_text(text)
    except Exception as e:
        st.error(f"Error reading Word file: {e}")
        return []

def create_chart(df, ticker):
    fig = make_subplots(
        rows=2, cols=1,
        shared_xaxes=True,
        vertical_spacing=0.08,
        row_heights=[0.7, 0.3],
        subplot_titles=(None, None)
    )

    fig.add_trace(
        go.Candlestick(
            x=df.index,
            open=df['Open'],
            high=df['High'],
            low=df['Low'],
            close=df['Close'],
            name='Price',
            increasing_line_color='#22c55e',
            decreasing_line_color='#ef4444',
            increasing_fillcolor='#22c55e',
            decreasing_fillcolor='#ef4444'
        ),
        row=1, col=1
    )

    fig.add_trace(
        go.Scatter(
            x=df.index,
            y=df['MA150'],
            mode='lines',
            name='150-Day MA',
            line=dict(color='#8b5cf6', width=3)
        ),
        row=1, col=1
    )

    fig.add_trace(
        go.Scatter(
            x=df.index,
            y=df['ATR'],
            mode='lines',
            name='ATR (14)',
            fill='tozeroy',
            line=dict(color='#3b82f6', width=2),
            fillcolor='rgba(59, 130, 246, 0.15)'
        ),
        row=2, col=1
    )

    fig.update_layout(
        height=550,
        template='plotly_white',
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="center",
            x=0.5,
            font=dict(size=14, color='#475569')
        ),
        xaxis_rangeslider_visible=False,
        margin=dict(l=20, r=20, t=60, b=20),
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        font=dict(family="Plus Jakarta Sans", color='#475569')
    )

    fig.update_xaxes(showgrid=True, gridwidth=1, gridcolor='rgba(0,0,0,0.05)', showline=False)
    fig.update_yaxes(showgrid=True, gridwidth=1, gridcolor='rgba(0,0,0,0.05)', showline=False)

    fig.add_annotation(
        text=f"<b>{ticker}</b> Price & 150-Day MA",
        xref="paper", yref="paper",
        x=0.5, y=1.12,
        showarrow=False,
        font=dict(size=20, color='#0f172a')
    )

    fig.add_annotation(
        text="<b>ATR (14-Day)</b>",
        xref="paper", yref="paper",
        x=0.5, y=0.28,
        showarrow=False,
        font=dict(size=16, color='#64748b')
    )

    return fig

# ============ MAIN APP ============

st.markdown('<h1 class="main-title">Stock <span>Analyzer</span> Pro</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-title">Real-time analysis with 150-Day Moving Average & ATR</p>', unsafe_allow_html=True)

tab1, tab2 = st.tabs(["📊 Single Stock", "📁 Portfolio"])

with tab1:
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        ticker = st.text_input(
            "",
            placeholder="AAPL",
            key="single_ticker",
            label_visibility="collapsed"
        ).upper().strip()

    if ticker:
        with st.spinner(f"Loading {ticker}..."):
            df, info = get_stock_data(ticker)

        if df is not None and not df.empty:
            current_price = df['Close'].iloc[-1]
            ma150 = df['MA150'].iloc[-1]
            atr = df['ATR'].iloc[-1]

            company_name = info.get('longName', ticker) if info else ticker
            st.markdown(f'<div class="company-name">{company_name}</div>', unsafe_allow_html=True)

            col1, col2, col3 = st.columns(3)

            with col1:
                price_change = ((current_price - df['Close'].iloc[-2]) / df['Close'].iloc[-2]) * 100
                change_class = "metric-change-up" if price_change >= 0 else "metric-change-down"
                change_symbol = "↑" if price_change >= 0 else "↓"
                st.markdown(f"""
                <div class="metric-card metric-card-price">
                    <div class="metric-label">Current Price</div>
                    <div class="metric-value metric-value-price">${current_price:.2f}</div>
                    <div class="metric-change {change_class}">{change_symbol} {abs(price_change):.2f}%</div>
                </div>
                """, unsafe_allow_html=True)

            with col2:
                ma_diff = ((current_price - ma150) / ma150) * 100 if pd.notna(ma150) else 0
                trend_class = "metric-change-up" if ma_diff >= 0 else "metric-change-down"
                trend_word = "Above" if ma_diff >= 0 else "Below"
                trend_symbol = "↑" if ma_diff >= 0 else "↓"
                st.markdown(f"""
                <div class="metric-card metric-card-ma">
                    <div class="metric-label">150-Day MA</div>
                    <div class="metric-value metric-value-ma">${ma150:.2f}</div>
                    <div class="metric-change {trend_class}">{trend_symbol} {trend_word} by {abs(ma_diff):.1f}%</div>
                </div>
                """, unsafe_allow_html=True)

            with col3:
                atr_pct = (atr / current_price) * 100 if current_price > 0 else 0
                st.markdown(f"""
                <div class="metric-card metric-card-atr">
                    <div class="metric-label">ATR (14-Day)</div>
                    <div class="metric-value metric-value-atr">${atr:.2f}</div>
                    <div class="metric-change" style="background: #fef3c7; color: #d97706;">{atr_pct:.2f}% volatility</div>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)

            fig = create_chart(df, ticker)
            st.plotly_chart(fig, use_container_width=True)

            with st.expander("📊 Key Statistics"):
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    high_52 = info.get('fiftyTwoWeekHigh') if info else None
                    st.metric("52-Week High", f"${high_52:.2f}" if high_52 else "N/A")
                with col2:
                    low_52 = info.get('fiftyTwoWeekLow') if info else None
                    st.metric("52-Week Low", f"${low_52:.2f}" if low_52 else "N/A")
                with col3:
                    st.metric("Volume", f"{df['Volume'].iloc[-1]:,.0f}")
                with col4:
                    avg_vol = info.get('averageVolume') if info else None
                    st.metric("Avg Volume", f"{avg_vol:,.0f}" if avg_vol else "N/A")
        else:
            st.error(f"❌ Could not find data for '{ticker}'. Please check the symbol.")
    else:
        st.markdown("<br>", unsafe_allow_html=True)
        st.info("👆 Enter a stock ticker symbol above to get started")

        st.markdown('<div class="section-header">Popular Tickers</div>', unsafe_allow_html=True)
        cols = st.columns(5)
        examples = ['AAPL', 'MSFT', 'GOOGL', 'AMZN', 'TSLA']
        for col, ex in zip(cols, examples):
            with col:
                st.markdown(f'<div class="ticker-chip">{ex}</div>', unsafe_allow_html=True)

with tab2:
    st.markdown('<div class="section-header">📁 Upload Your Portfolio</div>', unsafe_allow_html=True)
    st.markdown("Upload a file with your stock tickers to analyze your entire portfolio")

    st.markdown("<br>", unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "Drop your file here",
        type=['xlsx', 'xls', 'docx', 'pdf'],
        help="Supported: Excel, Word, PDF",
        label_visibility="collapsed"
    )

    if uploaded_file is not None:
        file_type = uploaded_file.name.split('.')[-1].lower()

        with st.spinner("Extracting tickers..."):
            if file_type in ['xlsx', 'xls']:
                tickers = read_excel_tickers(uploaded_file)
            elif file_type == 'docx':
                tickers = read_word_tickers(uploaded_file)
            elif file_type == 'pdf':
                tickers = read_pdf_tickers(uploaded_file)
            else:
                tickers = []
                st.error("Unsupported file format")

        if tickers:
            st.success(f"✓ Found {len(tickers)} tickers: {', '.join(tickers)}")

            edited_tickers = st.text_area(
                "Edit tickers if needed",
                value=", ".join(tickers),
                help="Remove false positives or add missing tickers"
            )

            if ',' in edited_tickers:
                final_tickers = [t.strip().upper() for t in edited_tickers.split(',') if t.strip()]
            else:
                final_tickers = [t.strip().upper() for t in edited_tickers.split('\n') if t.strip()]

            st.markdown("<br>", unsafe_allow_html=True)

            if st.button("🔍 Analyze Portfolio", type="primary", use_container_width=True):
                if final_tickers:
                    st.markdown('<div class="section-header">📊 Portfolio Analysis</div>', unsafe_allow_html=True)

                    progress_bar = st.progress(0)
                    status_text = st.empty()

                    results = []
                    failed_tickers = []

                    for i, ticker in enumerate(final_tickers):
                        status_text.text(f"Analyzing {ticker}...")
                        progress_bar.progress((i + 1) / len(final_tickers))

                        result = get_stock_summary(ticker)
                        if result:
                            results.append(result)
                        else:
                            failed_tickers.append(ticker)

                    status_text.empty()
                    progress_bar.empty()

                    if failed_tickers:
                        st.warning(f"⚠️ Could not fetch: {', '.join(failed_tickers)}")

                    if results:
                        df_results = pd.DataFrame(results)

                        df_display = df_results.copy()
                        df_display['Current Price'] = df_display['Current Price'].apply(lambda x: f"${x:.2f}" if pd.notna(x) else "N/A")
                        df_display['150-Day MA'] = df_display['150-Day MA'].apply(lambda x: f"${x:.2f}" if pd.notna(x) else "N/A")
                        df_display['ATR (14)'] = df_display['ATR (14)'].apply(lambda x: f"${x:.2f}" if pd.notna(x) else "N/A")

                        def color_gap(val):
                            if val == "N/A":
                                return val
                            num = float(val.replace('%', '').replace('+', ''))
                            if num >= 0:
                                return f"🟢 +{abs(num):.2f}%"
                            else:
                                return f"🔴 {num:.2f}%"

                        df_display['Gap %'] = df_results['Gap %'].apply(
                            lambda x: color_gap(f"{x:.2f}%") if pd.notna(x) else "N/A"
                        )

                        st.dataframe(
                            df_display,
                            use_container_width=True,
                            hide_index=True,
                            column_config={
                                "Ticker": st.column_config.TextColumn("Ticker", width="small"),
                                "Current Price": st.column_config.TextColumn("Price", width="medium"),
                                "150-Day MA": st.column_config.TextColumn("150-MA", width="medium"),
                                "Gap %": st.column_config.TextColumn("Gap %", width="medium"),
                                "ATR (14)": st.column_config.TextColumn("ATR", width="medium"),
                            }
                        )

                        st.markdown('<div class="section-header">📈 Summary</div>', unsafe_allow_html=True)

                        valid_gaps = df_results['Gap %'].dropna()
                        above_ma = (valid_gaps > 0).sum()
                        below_ma = (valid_gaps < 0).sum()
                        avg_gap = valid_gaps.mean()

                        col1, col2, col3 = st.columns(3)

                        with col1:
                            st.markdown(f"""
                            <div class="summary-card">
                                <div class="summary-number">{above_ma}</div>
                                <div class="summary-label">Above 150-MA</div>
                            </div>
                            """, unsafe_allow_html=True)

                        with col2:
                            st.markdown(f"""
                            <div class="summary-card">
                                <div class="summary-number">{below_ma}</div>
                                <div class="summary-label">Below 150-MA</div>
                            </div>
                            """, unsafe_allow_html=True)

                        with col3:
                            st.markdown(f"""
                            <div class="summary-card">
                                <div class="summary-number">{avg_gap:+.1f}%</div>
                                <div class="summary-label">Avg Gap</div>
                            </div>
                            """, unsafe_allow_html=True)

                        st.markdown("<br>", unsafe_allow_html=True)

                        csv = df_results.to_csv(index=False)
                        st.download_button(
                            label="📥 Download CSV",
                            data=csv,
                            file_name="portfolio_analysis.csv",
                            mime="text/csv",
                            use_container_width=True
                        )
                    else:
                        st.error("❌ Could not fetch data for any tickers.")
                else:
                    st.warning("Please enter at least one ticker.")
        else:
            st.warning("No tickers found in file. Check the content.")
    else:
        st.markdown("""
        <div class="upload-area">
            <div class="upload-icon">📄</div>
            <div class="upload-text">Drop your file here</div>
            <div class="upload-subtext">Supports Excel, Word & PDF</div>
        </div>
        """, unsafe_allow_html=True)

st.markdown("""
<div class="footer">
    Data provided by Yahoo Finance • ATR = Average True Range (14-day)
</div>
""", unsafe_allow_html=True)
